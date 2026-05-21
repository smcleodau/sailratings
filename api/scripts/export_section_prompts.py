"""Export every section's Claude prompt to a .docx Stuart can review.

For each section: title, the Facts fields Claude sees, the exact prompt
text. Order matches the report (s01 → s11).
"""

from __future__ import annotations

import sys
from dataclasses import fields
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches

from irc_data.api.services.report import prompts, facts as facts_mod

OUT = Path("/tmp/sailratings_section_prompts.docx")

NAVY = RGBColor(0x0A, 0x22, 0x40)
BRASS = RGBColor(0xC2, 0x9B, 0x61)
MUTED = RGBColor(0x64, 0x74, 0x8B)

# Section order on the page, mapped to its Facts dataclass + prompt
# constant in prompts.py. SYSTEM_PROMPT_V2 is the same shared prompt for
# every section.
SECTIONS = [
    ("s01", "Executive Summary",  "ExecutiveSummaryFacts",   "EXECUTIVE_SUMMARY_PROMPT"),
    ("s02", "Identity & History", "IdentityFacts",           "IDENTITY_PROMPT"),
    ("s03", "Rating Anatomy",     "RatingAnatomyFacts",      "RATING_ANATOMY_PROMPT"),
    ("s04", "Rating Evolution",   "RatingEvolutionFacts",    "RATING_EVOLUTION_PROMPT"),
    ("s05", "Class Context",      "ClassContextFacts",       "CLASS_CONTEXT_PROMPT"),
    ("s06", "Racing Performance", "PerformanceFacts",        "PERFORMANCE_PROMPT"),
    ("s07", "Measurement Sensitivity", "SensitivityFacts",   "SENSITIVITY_PROMPT"),
    ("s08", "Optimisation Recommendations", "OptimisationFacts", "OPTIMISATION_PROMPT"),
    ("s09", "Formula Drift",      "FormulaDriftFacts",       "FORMULA_DRIFT_PROMPT"),
    ("s10", "Rival Watch",        "RivalsFacts",             "RIVALS_PROMPT"),
    ("s11", "Appendix",           "AppendixFacts",           None),  # deterministic, no LLM
]


def _set_run(run, size=11, bold=False, color=None, italic=False):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _add_heading(doc, text, *, level=1, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    sizes = {1: 20, 2: 14, 3: 11}
    r = p.add_run(text)
    _set_run(r, size=sizes.get(level, 11), bold=True, color=color)


def _add_para(doc, text, *, size=11, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    _set_run(r, size=size, italic=italic, color=color)
    return p


def _add_mono_block(doc, text):
    """Render a code block — small monospace font."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _facts_fields_for(dataclass_name: str) -> list[tuple[str, str]]:
    cls = getattr(facts_mod, dataclass_name, None)
    if cls is None:
        return []
    out: list[tuple[str, str]] = []
    for f in fields(cls):
        # Walk one level into list[Identity] / list[RaceResultLite] etc.
        type_str = str(f.type).replace("typing.", "")
        out.append((f.name, type_str))
    return out


def main() -> int:
    doc = Document()

    # ── Cover ────────────────────────────────────────────────────────
    _add_heading(doc, "SailRatings Premium Report", level=1)
    _add_para(doc, "Section prompts — for review and editing", italic=True, color=MUTED)
    _add_para(doc, "")
    _add_para(
        doc,
        "Every premium report is built by running 11 section generators in parallel. "
        "Each section receives a strict Facts payload (what Claude can see) and a "
        "section-specific prompt (what Claude is told to write). The shared system "
        "prompt below sits in front of every section's user-prompt and enforces the "
        "truth-discipline rules — never cite a number not in Facts, never invent "
        "names, never speculate about the IRC formula.",
    )

    # ── System prompt ────────────────────────────────────────────────
    _add_heading(doc, "Shared system prompt (every section)", level=1, color=BRASS)
    _add_para(
        doc,
        "Sits in front of every section's user-prompt. Sets the voice, "
        "the absolute rules, and the style guide.",
        italic=True, color=MUTED,
    )
    _add_mono_block(doc, prompts.SYSTEM_PROMPT_V2.strip())

    # ── Per-section ──────────────────────────────────────────────────
    for s_id, title, facts_cls, prompt_name in SECTIONS:
        doc.add_page_break()
        _add_heading(doc, f"§{int(s_id[1:])}. {title}  ({s_id})", level=1)

        # Facts the section gets
        _add_heading(doc, "Facts payload (what Claude sees)", level=2, color=BRASS)
        ff = _facts_fields_for(facts_cls)
        if not ff:
            _add_para(doc, f"(no facts dataclass {facts_cls} found)", italic=True, color=MUTED)
        else:
            _add_para(
                doc,
                f"Source dataclass: {facts_cls}  ·  {len(ff)} fields",
                italic=True, color=MUTED,
            )
            mono = "\n".join(f"  {name:<25}  {typ}" for name, typ in ff)
            _add_mono_block(doc, mono)

        # The prompt itself
        _add_heading(doc, "Section prompt", level=2, color=BRASS)
        if prompt_name is None:
            _add_para(
                doc,
                "Deterministic section — no Claude call. The Appendix is rendered "
                "straight from the AppendixFacts (methodology blurb, six data "
                "sources, twelve-term glossary).",
                italic=True, color=MUTED,
            )
        else:
            prompt_text = getattr(prompts, prompt_name, None)
            if prompt_text is None:
                _add_para(doc, f"(prompt {prompt_name} not found)", italic=True, color=MUTED)
            else:
                _add_para(
                    doc, f"Constant name: {prompt_name}",
                    italic=True, color=MUTED,
                )
                _add_mono_block(doc, prompt_text.strip())

    doc.save(OUT)
    print(f"Saved → {OUT}  ({OUT.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
